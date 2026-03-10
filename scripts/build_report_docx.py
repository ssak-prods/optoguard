from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(r"d:\optoguard")
TEXT_PATH = ROOT / "technical-report.txt"
PLOTS = [
    ("Figure 1. Uncertainty vs. lighting condition", ROOT / "plotsfolder" / "uncVLightV2.png"),
    ("Figure 2. Uncertainty vs. occlusion level", ROOT / "plotsfolder" / "uncVoccV2.png"),
    ("Figure 3. Latency by condition and hardware", ROOT / "plotsfolder" / "latV2.png"),
    ("Figure 4. Development timeline (Oct 2025–Mar 2026)", ROOT / "plotsfolder" / "timeline.png"),
]

OUT_DIR = ROOT / "paper"
OUT_PATH = OUT_DIR / "OptoGuard_Technical_Report.docx"


def add_field_toc(paragraph, instruction: str) -> None:
    """Insert a field code (e.g. TOC) that Word can update to show page numbers."""
    # Adapted from python-docx examples for adding TOC fields.
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    run._r.append(instr_text)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_separate)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    raw = TEXT_PATH.read_text(encoding="utf-8")
    lines = [ln.rstrip() for ln in raw.splitlines()]

    doc = Document()

    # Front matter: contents and list of figures (Word will populate page numbers after updating fields).
    doc.add_heading("Contents", level=1)
    p_toc = doc.add_paragraph()
    add_field_toc(p_toc, 'TOC \\o "1-2" \\h \\z \\u')

    doc.add_page_break()

    doc.add_heading("List of Figures", level=1)
    p_lof = doc.add_paragraph()
    # This relies on figure captions being styled as 'Caption' later in the document.
    add_field_toc(p_lof, 'TOC \\h \\z \\c "Figure"')

    doc.add_page_break()

    # Main title and body.
    doc.add_heading(
        "OptoGuard: Uncertainty-Aware Object Detection under Distribution Shift on Edge Hardware",
        level=0,
    )
    doc.add_paragraph(
        "Technical report on Monte Carlo Dropout-based uncertainty analysis for YOLOv8n under distribution shift."
    )
    doc.add_paragraph("")

    for ln in lines:
        s = ln.strip()
        if not s or set(s) == {"_"}:
            continue

        # Top-level headings such as "1. Introduction"
        if len(s) > 3 and s[0].isdigit() and s[1:3] == ". ":
            doc.add_heading(s, level=1)
            continue

        # Sub-headings such as "3.1 Base Detector ..."
        parts = s.split(" ", 1)
        if parts and "." in parts[0] and parts[0].replace(".", "").isdigit() and len(parts[0]) >= 3:
            doc.add_heading(s, level=2)
            continue

        doc.add_paragraph(s)

    # Figures with caption style so that the List of Figures field can pick them up.
    doc.add_page_break()
    doc.add_heading("Figures", level=1)
    for caption, path in PLOTS:
        if path.exists():
            cap_para = doc.add_paragraph(caption)
            try:
                cap_para.style = "Caption"
            except Exception:
                # If the template has no 'Caption' style, fall back silently.
                pass
            doc.add_picture(str(path), width=Inches(6.5))
            doc.add_paragraph("")

    doc.save(str(OUT_PATH))
    print(OUT_PATH)


if __name__ == "__main__":
    main()

